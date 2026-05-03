import re
import io
import logging
import time
import subprocess
import tempfile
import os
from typing import List, Dict, Any, Optional, Set
import pdfplumber
from docx import Document

from app.backend.services.metrics import RESUME_PARSE_DURATION

logger = logging.getLogger(__name__)

try:
    import fitz as pymupdf   # PyMuPDF
    _HAS_PYMUPDF = True
except ImportError:
    _HAS_PYMUPDF = False

try:
    from unidecode import unidecode as _unidecode
    _HAS_UNIDECODE = True
except ImportError:
    _HAS_UNIDECODE = False

try:
    import docx2txt
    _HAS_DOCX2TXT = True
except ImportError:
    _HAS_DOCX2TXT = False

try:
    from striprtf.striprtf import rtf_to_text
    _HAS_STRIPRTF = True
except ImportError:
    _HAS_STRIPRTF = False

try:
    from odf import opendocument
    from odf.text import P
    _HAS_ODFPY = True
except ImportError:
    _HAS_ODFPY = False

try:
    import dateparser
    _HAS_DATEPARSER = True
except ImportError:
    _HAS_DATEPARSER = False


# ─── spaCy NER singleton for name extraction (Tier 0) ────────────────────────

_spacy_nlp = None

def _get_spacy_model():
    """Lazy-load spaCy model once; return None if unavailable."""
    global _spacy_nlp
    if _spacy_nlp is None:
        try:
            import spacy
            _spacy_nlp = spacy.load("en_core_web_sm")
        except (ImportError, OSError):
            _spacy_nlp = False  # Mark as unavailable to avoid repeated attempts
    return _spacy_nlp if _spacy_nlp is not False else None


def _extract_name_ner(raw_text: str) -> str | None:
    """Tier 0: Extract candidate name using spaCy NER on first 50 lines."""
    nlp = _get_spacy_model()
    if nlp is None:
        return None
    
    # Only process first 50 lines (header area of resume)
    lines = raw_text.strip().split('\n')[:50]
    header_text = '\n'.join(lines)
    
    doc = nlp(header_text)
    
    # Find first PERSON entity
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            # Basic validation: 1-5 words, no digits, reasonable length
            words = name.split()
            if 1 <= len(words) <= 5 and not any(c.isdigit() for c in name) and len(name) <= 60:
                return name
    
    return None


# ─── Date extraction helpers (dateparser-powered) ────────────────────────────

_PRESENT_SYNONYMS_RE = re.compile(
    r'\b(till\s*date|till\s*now|till\s*present|to\s*date|to\s*present'
    r'|ongoing|continuing|current\s*position)\b', re.IGNORECASE)

_MONTH_PERIOD_RE = re.compile(
    r'\b(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|SEPT|OCT|NOV|DEC)\.', re.IGNORECASE)


def _normalize_date_text(text: str) -> str:
    """Normalize common date formatting variations before matching."""
    s = _MONTH_PERIOD_RE.sub(r'\1', text)       # "AUG." -> "AUG"
    s = _PRESENT_SYNONYMS_RE.sub('Present', s)  # "Till Date" -> "Present"
    return s


_DATE_RANGE_CANDIDATE_RE = re.compile(
    r'('
    r'(?:\d{1,2}[/\-]\d{4})'           # MM/YYYY or MM-YYYY
    r'|(?:\d{4}[/\-]\d{1,2})'           # YYYY/MM or YYYY-MM
    r'|(?:[A-Za-z]+\.?\s+\d{4})'        # Month YYYY (with optional period)
    r'|(?:\d{4})'                         # Bare year
    r')'
    r'\s*[-–—]+\s*'                       # Separator (dash variants)
    r'('
    r'(?:present|current|now)'
    r'|(?:\d{1,2}[/\-]\d{4})'
    r'|(?:\d{4}[/\-]\d{1,2})'
    r'|(?:[A-Za-z]+\.?\s+\d{4})'
    r'|(?:\d{4})'
    r')',
    re.IGNORECASE
)


def _parse_date_str(s: str):
    """Parse a date string using dateparser. Returns datetime or None."""
    if not s:
        return None
    s = s.strip()
    if s.lower() in ('present', 'current', 'now'):
        return None  # Caller handles "present" specially
    if not _HAS_DATEPARSER:
        return None
    dt = dateparser.parse(s, settings={
        'PREFER_DAY_OF_MONTH': 'first',
        'REQUIRE_PARTS': ['year'],
    })
    return dt


def _extract_date_range(line: str):
    """Extract a date range from a line using normalization + dateparser.
    Returns (start_str, end_str) or None.
    start_str/end_str are normalized like "August 2011" or "present".
    """
    normalized = _normalize_date_text(line)
    m = _DATE_RANGE_CANDIDATE_RE.search(normalized)
    if not m:
        # Year-only fallback: look for two bare years
        years = re.findall(r'\b((?:19|20)\d{2})\b', normalized)
        if len(years) >= 2:
            return (years[0], years[-1])
        if len(years) == 1 and re.search(r'\b(present|current|now)\b', normalized, re.IGNORECASE):
            return (years[0], 'present')
        return None

    raw_start, raw_end = m.group(1).strip(), m.group(2).strip()

    if raw_end.lower() in ('present', 'current', 'now'):
        end_str = 'present'
    else:
        dt_end = _parse_date_str(raw_end)
        end_str = dt_end.strftime('%B %Y') if dt_end else raw_end

    dt_start = _parse_date_str(raw_start)
    start_str = dt_start.strftime('%B %Y') if dt_start else raw_start

    return (start_str, end_str)


# ─── JD text extractor (multi-format, lenient) ───────────────────────────────

def extract_jd_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from a Job Description file.

    Supports: PDF, DOCX, DOC (legacy binary), TXT, RTF, HTML/HTM, ODT,
    Markdown, and any plain-text file regardless of extension.

    Unlike parse_resume(), this function only needs raw text — it does NOT
    return structured data and is intentionally permissive about format.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # ── PDF ──────────────────────────────────────────────────────────────────
    if ext == "pdf":
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            if text.strip():
                return text
        except Exception as e:
            logger.warning("Non-critical: PDF extraction failed for %s, falling back to plain text: %s", filename, e)
    
    # ── DOCX (modern Word .docx / Office Open XML) ───────────────────────────
    if ext == "docx":
        try:
            doc = Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
            if text.strip():
                return text
        except Exception as e:
            logger.warning("Non-critical: DOCX extraction failed for %s, falling back to plain text: %s", filename, e)
    
    # ── DOC (legacy binary Word) — best-effort ASCII extraction ──────────────
    if ext == "doc":
        try:
            raw = file_bytes.decode("latin-1", errors="ignore")
            # Remove non-printable characters except newline/tab
            cleaned = re.sub(r"[^\x20-\x7E\n\t]", " ", raw)
            # Collapse multiple spaces but keep newlines
            cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
            lines = [l.strip() for l in cleaned.splitlines() if len(l.strip()) > 4]
            text = "\n".join(lines)
            if text.strip():
                return text
        except Exception as e:
            logger.warning("Non-critical: DOC extraction failed for %s, falling back to plain text: %s", filename, e)
    
    # ── RTF ───────────────────────────────────────────────────────────────────
    if ext == "rtf":
        try:
            raw = file_bytes.decode("latin-1", errors="ignore")
            # Strip RTF control words, groups, and backslash escapes
            text = re.sub(r"\\[a-z]+\-\d*\s?", " ", raw)
            text = re.sub(r"\{[^{}]*\}", " ", text)
            text = re.sub(r"[{}\\]", " ", text)
            text = re.sub(r"\s+", " ", text)
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning("Non-critical: RTF extraction failed for %s, falling back to plain text: %s", filename, e)
    
    # ── HTML / HTM ────────────────────────────────────────────────────────────
    if ext in ("html", "htm"):
        try:
            raw = file_bytes.decode("utf-8", errors="ignore")
            text = re.sub(r"<[^>]+>", " ", raw)
            text = re.sub(r"\s+", " ", text)
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning("Non-critical: HTML extraction failed for %s, falling back to plain text: %s", filename, e)
    
    # ── ODT (Open Document Text) — it's a ZIP; extract content.xml ───────────
    if ext == "odt":
        try:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                with z.open("content.xml") as f:
                    xml = f.read().decode("utf-8", errors="ignore")
            text = re.sub(r"<[^>]+>", " ", xml)
            text = re.sub(r"\s+", " ", text)
            if text.strip():
                return text.strip()
        except Exception as e:
            logger.warning("Non-critical: ODT extraction failed for %s, falling back to plain text: %s", filename, e)

    # ── Plain text fallback (TXT, MD, CSV, unknown, or any failed attempt) ───
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            text = file_bytes.decode(encoding)
            if text.strip():
                return text
        except (UnicodeDecodeError, LookupError):
            continue

    raise ValueError(
        f"Could not extract readable text from '{filename}'. "
        "Supported formats: PDF, DOCX, DOC, TXT, RTF, HTML, ODT."
    )


class ResumeParser:
    def __init__(self):
        # Order: more specific / common CV formats first (year-to-present is very common).
        self.date_patterns = [
            r'(\d{4})\s*[-–—]\s*(present|current|now)',
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(present|current|now)',
            r'(\w+\s+\d{4})\s*[-–—]\s*(present|current|now)',
            r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|present|current|now)',
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4}|present|current|now)',
            r'(\d{4})\s*[-–—]\s*(\d{4}|present|current|now)',
        ]

    # Supported formats mapping
    SUPPORTED_FORMATS = {
        '.pdf': '_extract_pdf_multistage',
        '.docx': '_extract_docx_multistage',
        '.doc': '_extract_doc_multistage',
        '.txt': '_extract_txt',
        '.rtf': '_extract_rtf',
        '.odt': '_extract_odt',
    }

    def extract_text(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from resume file using multi-stage pipeline."""
        ext = os.path.splitext(filename.lower())[1]
        
        if ext == '.pdf':
            return self._extract_pdf_multistage(file_bytes)
        elif ext == '.docx':
            return self._extract_docx_multistage(file_bytes)
        elif ext == '.doc':
            return self._extract_doc_multistage(file_bytes)
        elif ext == '.txt':
            return self._extract_txt(file_bytes)
        elif ext == '.rtf':
            return self._extract_rtf(file_bytes)
        elif ext == '.odt':
            return self._extract_odt(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {filename}. Supported: {list(self.SUPPORTED_FORMATS.keys())}")

    def _extract_pdf_multistage(self, file_bytes: bytes) -> str:
        """
        Multi-stage PDF extraction pipeline.
        
        Stage 1: PyMuPDF (best for multi-column, correct reading order)
        Stage 2: pdfplumber with table extraction (best for structured tables)
        Stage 3: Combined deduplicated output
        """
        text_parts = []
        seen_texts: Set[str] = set()
        
        def add_unique_text(text: str, source: str) -> None:
            """Add text if not already seen (deduplication)."""
            if not text or not text.strip():
                return
            # Normalize for comparison
            normalized = re.sub(r'\s+', ' ', text.strip().lower())
            if normalized and normalized not in seen_texts and len(normalized) > 3:
                seen_texts.add(normalized)
                text_parts.append((text.strip(), source))
        
        # Stage 1: PyMuPDF extraction (primary - best for layout preservation)
        if _HAS_PYMUPDF:
            try:
                doc = pymupdf.open(stream=file_bytes, filetype="pdf")
                for page_num, page in enumerate(doc):
                    page_text = page.get_text("text")
                    if page_text.strip():
                        add_unique_text(page_text, f"pymupdf_page_{page_num}")
                doc.close()
                logger.info("PyMuPDF extraction completed")
            except Exception as e:
                logger.warning("PyMuPDF extraction failed: %s", e)
        
        # Stage 2: pdfplumber with table extraction (for tables and structured data)
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract tables first (contact info often in tables)
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) for cell in row if cell)
                                if row_text.strip():
                                    add_unique_text(row_text, f"table_{page_num}_{table_idx}")
                    
                    # Extract regular text
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        add_unique_text(page_text, f"pdfplumber_page_{page_num}")
            logger.info("pdfplumber extraction completed with table support")
        except Exception as e:
            logger.warning("pdfplumber extraction failed: %s", e)
        
        # Combine all extracted text with source priority
        # Prioritize table content (often contains contact info) and header areas
        combined_parts = []
        table_parts = []
        text_parts_list = []
        
        for text, source in text_parts:
            if 'table' in source:
                table_parts.append(text)
            else:
                text_parts_list.append(text)
        
        # Order: tables first (contact info), then regular text
        combined_parts = table_parts + text_parts_list
        
        # Final deduplication pass
        final_lines = []
        final_seen: Set[str] = set()
        for part in combined_parts:
            for line in part.split('\n'):
                line = line.strip()
                if not line:
                    continue
                normalized = re.sub(r'\s+', ' ', line.lower())
                if normalized not in final_seen and len(line) > 1:
                    final_seen.add(normalized)
                    final_lines.append(line)
        
        text = '\n'.join(final_lines)
        
        # Scanned-PDF guard
        if len(text.strip()) < 100:
            raise ValueError(
                "This PDF appears to be a scanned image and cannot be read automatically. "
                "Please upload a text-based PDF (exported from Word/Google Docs) rather than "
                "a scanned or photographed document."
            )
        
        # Normalise Unicode characters
        if _HAS_UNIDECODE:
            text = _unidecode(text)
        
        logger.info(f"Multi-stage PDF extraction: extracted {len(final_lines)} unique lines")
        return text

    def _extract_pdf(self, file_bytes: bytes) -> str:
        """Legacy PDF extraction - kept for backward compatibility."""
        return self._extract_pdf_multistage(file_bytes)

    def _extract_docx_multistage(self, file_bytes: bytes) -> str:
        """
        Multi-stage DOCX extraction pipeline for maximum content recovery.
        
        Stage 1: Headers (contact info often in headers)
        Stage 2: Textboxes and shapes using docx2txt
        Stage 3: Tables (contact info often in tables)
        Stage 4: Regular paragraphs
        Stage 5: XML fallback for corrupted files
        """
        text_parts = []
        seen_texts: Set[str] = set()
        
        def add_unique_text(text: str, source: str) -> None:
            """Add text if not already seen (deduplication)."""
            if not text or not text.strip():
                return
            # Normalize for comparison
            normalized = re.sub(r'\s+', ' ', text.strip().lower())
            if normalized and normalized not in seen_texts and len(normalized) > 2:
                seen_texts.add(normalized)
                text_parts.append((text.strip(), source))
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            
            # Stage 1: Extract text from headers (often contains contact info)
            try:
                for section_idx, section in enumerate(doc.sections):
                    header = section.header
                    header_texts = []
                    for para in header.paragraphs:
                        if para.text.strip():
                            header_texts.append(para.text.strip())
                    if header_texts:
                        header_text = "\n".join(header_texts)
                        add_unique_text(header_text, f"header_{section_idx}")
                logger.info("DOCX header extraction completed")
            except Exception as e:
                logger.warning("Header extraction failed: %s", e)
            
            # Stage 2: Extract text from textboxes/shapes using docx2txt
            if _HAS_DOCX2TXT:
                try:
                    # Save to temp file for docx2txt
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    
                    try:
                        textbox_text = docx2txt.process(tmp_path)
                        if textbox_text.strip():
                            add_unique_text(textbox_text, "textboxes_docx2txt")
                            logger.info("DOCX textbox extraction completed via docx2txt")
                    finally:
                        os.unlink(tmp_path)
                except Exception as e:
                    logger.warning("docx2txt extraction failed: %s", e)
            
            # Stage 3: Extract text from tables (contact info often in tables)
            try:
                for table_idx, table in enumerate(doc.tables):
                    for row_idx, row in enumerate(table.rows):
                        row_texts = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_texts.append(cell.text.strip())
                        if row_texts:
                            row_text = " | ".join(row_texts)
                            add_unique_text(row_text, f"table_{table_idx}_row_{row_idx}")
                logger.info("DOCX table extraction completed")
            except Exception as e:
                logger.warning("Table extraction failed: %s", e)
            
            # Stage 4: Extract text from paragraphs
            try:
                paragraph_texts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraph_texts.append(para.text.strip())
                if paragraph_texts:
                    add_unique_text("\n".join(paragraph_texts), "paragraphs")
                logger.info("DOCX paragraph extraction completed")
            except Exception as e:
                logger.warning("Paragraph extraction failed: %s", e)
                
        except Exception as e:
            logger.warning("python-docx extraction failed, trying XML fallback: %s", e)
        
        # Stage 5: XML fallback for corrupted/unusual files
        if not text_parts:
            try:
                import zipfile
                import xml.etree.ElementTree as ET
                
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx_zip:
                    # Try to find and read document.xml
                    if 'word/document.xml' in docx_zip.namelist():
                        xml_content = docx_zip.read('word/document.xml')
                        tree = ET.fromstring(xml_content)
                        
                        # Extract all text nodes
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        text_elements = tree.findall('.//w:t', ns)
                        xml_text = '\n'.join([elem.text for elem in text_elements if elem.text])
                        
                        if xml_text.strip():
                            add_unique_text(xml_text, "xml_fallback")
                            logger.info("DOCX XML fallback extraction completed")
            except Exception as e:
                logger.warning("XML fallback also failed: %s", e)
        
        if not text_parts:
            raise ValueError(
                "Unable to extract text from this Word document. "
                "The file may be corrupted or in an unsupported format. "
                "Please try re-saving the document or converting to PDF."
            )
        
        # Combine all extracted text with source priority
        # Order: headers first (contact info), then textboxes, then tables, then paragraphs
        source_priority = {
            'header': 0,
            'textbox': 1,
            'table': 2,
            'paragraph': 3,
            'xml': 4,
        }
        
        sorted_parts = sorted(text_parts, key=lambda x: (
            next((v for k, v in source_priority.items() if k in x[1]), 5),
            x[1]
        ))
        
        # Final deduplication pass
        final_lines = []
        final_seen: Set[str] = set()
        for text, source in sorted_parts:
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                normalized = re.sub(r'\s+', ' ', line.lower())
                if normalized not in final_seen and len(line) > 1:
                    final_seen.add(normalized)
                    final_lines.append(line)
        
        text = '\n'.join(final_lines)
        logger.info(f"Multi-stage DOCX extraction: extracted {len(final_lines)} unique lines")
        return text

    def _extract_docx(self, file_bytes: bytes) -> str:
        """Legacy DOCX extraction - kept for backward compatibility."""
        return self._extract_docx_multistage(file_bytes)

    def _extract_doc_multistage(self, file_bytes: bytes) -> str:
        """
        Multi-stage DOC (legacy Word) extraction pipeline.

        Stage 1: Try antiword if available (best for .doc files)
        Stage 2: Try LibreOffice headless conversion to PDF + pdfplumber
        Stage 3: Try LibreOffice headless conversion to DOCX
        Stage 4: Best-effort ASCII extraction
        """
        text_parts = []

        # Stage 1: Try antiword if available
        try:
            result = subprocess.run(
                ['antiword', '-'],
                input=file_bytes,
                capture_output=True,
                timeout=10
            )
            if result.returncode == 0 and result.stdout:
                text = result.stdout.decode('utf-8', errors='ignore')
                if text.strip():
                    text_parts.append((text, "antiword"))
                    logger.info("DOC extraction completed via antiword")
        except (subprocess.SubprocessError, FileNotFoundError, Exception) as e:
            logger.debug("antiword not available or failed: %s", e)

        # Stage 2: Try LibreOffice PDF conversion + pdfplumber (best quality on Windows)
        if not text_parts:
            try:
                from app.backend.services.doc_converter import convert_to_pdf, extract_text_from_pdf_bytes
                pdf_bytes = convert_to_pdf(file_bytes, "resume.doc")
                if pdf_bytes:
                    text = extract_text_from_pdf_bytes(pdf_bytes)
                    if text.strip():
                        text_parts.append((text, "libreoffice_pdf"))
                        logger.info("DOC extraction completed via LibreOffice PDF conversion")
            except Exception as e:
                logger.debug("LibreOffice PDF conversion not available or failed: %s", e)

        # Stage 3: Try LibreOffice DOCX conversion
        if not text_parts:
            try:
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp_in:
                    tmp_in.write(file_bytes)
                    tmp_in_path = tmp_in.name

                with tempfile.TemporaryDirectory() as tmp_dir:
                    result = subprocess.run(
                        ['soffice', '--headless', '--convert-to', 'docx',
                         '--outdir', tmp_dir, tmp_in_path],
                        capture_output=True,
                        timeout=30
                    )

                    if result.returncode == 0:
                        converted_files = [f for f in os.listdir(tmp_dir) if f.endswith('.docx')]
                        if converted_files:
                            docx_path = os.path.join(tmp_dir, converted_files[0])
                            with open(docx_path, 'rb') as f:
                                docx_bytes = f.read()
                            text = self._extract_docx_multistage(docx_bytes)
                            if text.strip():
                                text_parts.append((text, "libreoffice_docx"))
                                logger.info("DOC extraction completed via LibreOffice DOCX conversion")

                os.unlink(tmp_in_path)
            except (subprocess.SubprocessError, FileNotFoundError, Exception) as e:
                logger.debug("LibreOffice DOCX conversion not available or failed: %s", e)

        # Stage 4: Best-effort ASCII extraction
        if not text_parts:
            try:
                raw = file_bytes.decode("latin-1", errors="ignore")
                cleaned = re.sub(r"[^\x20-\x7E\n\t]", " ", raw)
                cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
                lines = [l.strip() for l in cleaned.splitlines() if len(l.strip()) > 4]
                text = "\n".join(lines)
                if text.strip():
                    text_parts.append((text, "ascii_fallback"))
                    logger.info("DOC extraction completed via ASCII fallback")
            except Exception as e:
                logger.warning("ASCII fallback failed: %s", e)

        if not text_parts:
            raise ValueError(
                "Unable to extract text from this legacy Word document (.doc). "
                "Please convert to DOCX or PDF and try again."
            )

        # Combine and deduplicate
        final_lines = []
        final_seen: Set[str] = set()
        for text, source in text_parts:
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                normalized = re.sub(r'\s+', ' ', line.lower())
                if normalized not in final_seen and len(line) > 1:
                    final_seen.add(normalized)
                    final_lines.append(line)

        text = '\n'.join(final_lines)
        logger.info(f"Multi-stage DOC extraction: extracted {len(final_lines)} unique lines")
        return text

    def _extract_txt(self, file_bytes: bytes) -> str:
        """Extract text from TXT file with multiple encoding attempts."""
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                if text.strip():
                    logger.info(f"TXT extraction completed with {encoding} encoding")
                    return text
            except (UnicodeDecodeError, LookupError):
                continue
        
        raise ValueError("Unable to decode text file. Unsupported encoding.")

    def _extract_rtf(self, file_bytes: bytes) -> str:
        """
        Multi-stage RTF extraction pipeline.
        
        Stage 1: Use striprtf library if available
        Stage 2: Regex-based RTF stripping fallback
        """
        text_parts = []
        
        # Stage 1: Use striprtf library
        if _HAS_STRIPRTF:
            try:
                raw = file_bytes.decode('utf-8', errors='ignore')
                text = rtf_to_text(raw)
                if text.strip():
                    text_parts.append((text, "striprtf"))
                    logger.info("RTF extraction completed via striprtf")
            except Exception as e:
                logger.warning("striprtf extraction failed: %s", e)
        
        # Stage 2: Regex-based fallback
        if not text_parts:
            try:
                raw = file_bytes.decode("latin-1", errors="ignore")
                # Strip RTF control words, groups, and backslash escapes
                text = re.sub(r"\\[a-z]+\-?\d*\s?", " ", raw)
                text = re.sub(r"\{[^{}]*\}", " ", text)
                text = re.sub(r"[{}\\]", " ", text)
                text = re.sub(r"\s+", " ", text)
                if text.strip():
                    text_parts.append((text.strip(), "regex_fallback"))
                    logger.info("RTF extraction completed via regex fallback")
            except Exception as e:
                logger.warning("RTF regex fallback failed: %s", e)
        
        if not text_parts:
            raise ValueError("Unable to extract text from RTF file.")
        
        # Combine and deduplicate
        final_lines = []
        final_seen: Set[str] = set()
        for text, source in text_parts:
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                normalized = re.sub(r'\s+', ' ', line.lower())
                if normalized not in final_seen and len(line) > 1:
                    final_seen.add(normalized)
                    final_lines.append(line)
        
        text = '\n'.join(final_lines)
        logger.info(f"Multi-stage RTF extraction: extracted {len(final_lines)} unique lines")
        return text

    def _extract_odt(self, file_bytes: bytes) -> str:
        """
        Multi-stage ODT (OpenDocument Text) extraction pipeline.
        
        Stage 1: Use odfpy library if available
        Stage 2: ZIP-based XML extraction fallback
        """
        text_parts = []
        
        # Stage 1: Use odfpy library
        if _HAS_ODFPY:
            try:
                with tempfile.NamedTemporaryFile(suffix='.odt', delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                
                try:
                    doc = opendocument.load(tmp_path)
                    paragraphs = []
                    for paragraph in doc.getElementsByType(P):
                        text_content = str(paragraph)
                        # Extract text between tags
                        text_match = re.search(r'>([^<]+)<', text_content)
                        if text_match:
                            paragraphs.append(text_match.group(1))
                    
                    if paragraphs:
                        text = '\n'.join(paragraphs)
                        if text.strip():
                            text_parts.append((text, "odfpy"))
                            logger.info("ODT extraction completed via odfpy")
                finally:
                    os.unlink(tmp_path)
            except Exception as e:
                logger.warning("odfpy extraction failed: %s", e)
        
        # Stage 2: ZIP-based XML extraction
        if not text_parts:
            try:
                import zipfile
                
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                    with z.open("content.xml") as f:
                        xml = f.read().decode("utf-8", errors="ignore")
                
                text = re.sub(r"<[^>]+>", " ", xml)
                text = re.sub(r"\s+", " ", text)
                if text.strip():
                    text_parts.append((text.strip(), "xml_fallback"))
                    logger.info("ODT extraction completed via XML fallback")
            except Exception as e:
                logger.warning("ODT XML fallback failed: %s", e)
        
        if not text_parts:
            raise ValueError("Unable to extract text from ODT file.")
        
        # Combine and deduplicate
        final_lines = []
        final_seen: Set[str] = set()
        for text, source in text_parts:
            for line in text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                normalized = re.sub(r'\s+', ' ', line.lower())
                if normalized not in final_seen and len(line) > 1:
                    final_seen.add(normalized)
                    final_lines.append(line)
        
        text = '\n'.join(final_lines)
        logger.info(f"Multi-stage ODT extraction: extracted {len(final_lines)} unique lines")
        return text

    def parse_resume(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        text = self.extract_text(file_bytes, filename)

        return {
            "raw_text": text,
            "work_experience": self._extract_work_experience(text),
            "skills": self._extract_skills(text),
            "education": self._extract_education(text),
            "contact_info": self._extract_contact_info(text),
            "certifications": self._extract_certifications(text),
            "languages": self._extract_languages(text),
            "professional_summary": self._extract_professional_summary(text),
        }

    # ── Title/company splitting patterns ─────────────────────────────────
    _TITLE_COMPANY_DELIMITERS = re.compile(
        r'\s*(?:\||\t+|\s{3,}|\s+at\s+|\s+-\s+|\s+@\s+|,\s+(?=[A-Z]))\s*')

    _TITLE_KEYWORDS = re.compile(
        r'\b(Engineer|Developer|Manager|Director|Analyst|Architect|'
        r'Designer|Consultant|Administrator|Specialist|Coordinator|'
        r'Lead|Senior|Junior|Principal|Staff|VP|CTO|CEO|CFO|COO|'
        r'Intern|Trainee|Associate|Officer|Executive|Supervisor|'
        r'Technician|Programmer|Scientist|Researcher)\b',
        re.IGNORECASE)

    _COMPANY_KEYWORDS = re.compile(
        r'\b(Ltd\.?|Inc\.?|Corp\.?|LLC|LLP|Pvt\.?|Private|Limited|'
        r'Corporation|Company|Co\.?|Group|Solutions|Technologies|'
        r'Systems|Services|Software|Labs?|Laboratories|Industries|'
        r'Enterprises|Associates|Partners|International|Global)\b',
        re.IGNORECASE)

    def _split_title_company(self, text: str) -> tuple:
        """Split text into (title, company) using delimiters and keyword disambiguation.

        Handles patterns like:
        - "Senior Engineer | TechCorp"
        - "Senior Engineer at TechCorp"
        - "Senior Engineer - TechCorp"
        - "Senior Engineer @ TechCorp"
        - "Senior Engineer TRIGENT SOFTWARE LTD. deputed @ AMETEK"
        """
        text = text.strip()

        # Handle "deputed @" or "deputed at" pattern — extract host company
        deputed_match = re.search(
            r'\b(deputed\s+(?:@|at)\s+)(.+)', text, re.IGNORECASE)
        if deputed_match:
            host_company = deputed_match.group(2).strip()
            remainder = text[:deputed_match.start()].strip()
            # Remainder is "Title CONTRACTING_CO" or just "Title"
            if remainder:
                inner_title, inner_company = self._split_title_company(remainder)
                if inner_company:
                    return (inner_title, inner_company + " / " + host_company)
                return (inner_title, host_company)
            return ("", host_company)

        # Try splitting on delimiters
        parts = self._TITLE_COMPANY_DELIMITERS.split(text, maxsplit=1)

        if len(parts) >= 2:
            part1, part2 = parts[0].strip(), parts[1].strip()

            # Disambiguate: which part is title, which is company?
            p1_is_company = bool(self._COMPANY_KEYWORDS.search(part1))
            p2_is_company = bool(self._COMPANY_KEYWORDS.search(part2))
            p1_is_title = bool(self._TITLE_KEYWORDS.search(part1))
            p2_is_title = bool(self._TITLE_KEYWORDS.search(part2))

            if p1_is_company and not p2_is_company:
                return (part2, part1)  # company first, then title
            elif p1_is_title and not p2_is_title:
                return (part1, part2)  # title first, then company
            elif p2_is_company:
                return (part1, part2)  # assume title first
            else:
                return (part1, part2)  # default: title first

        # Single text — try to split at title/company boundary using keywords
        title_match = self._TITLE_KEYWORDS.search(text)
        company_match = self._COMPANY_KEYWORDS.search(text)

        if title_match and company_match and title_match.end() <= company_match.start():
            # Both present and title comes first: find the end of the title phrase.
            # Use the LAST title keyword match to find the split boundary.
            # For "Senior Engineer TRIGENT SOFTWARE LTD.", the last title keyword
            # is "Engineer" — we split after it, keeping "Senior Engineer" as title.
            last_title_end = title_match.end()
            for m in self._TITLE_KEYWORDS.finditer(text):
                if m.end() <= company_match.start():
                    last_title_end = m.end()
                else:
                    break

            # Now skip whitespace after the last title keyword to find the
            # start of the company portion.
            pos = last_title_end
            while pos < len(text) and text[pos] == ' ':
                pos += 1
            if pos < len(text):
                return (text[:last_title_end].strip(), text[pos:].strip())

        # Single text — try to determine if it's a title or company
        if self._TITLE_KEYWORDS.search(text):
            return (text, "")
        elif self._COMPANY_KEYWORDS.search(text):
            return ("", text)
        return (text, "")

    def _extract_work_experience(self, text: str) -> List[Dict[str, Any]]:
        jobs = []

        lines = text.split('\n')
        current_job = None
        # When dates sit on their own line ("2018 – Present"), company/title are often on the line above.
        last_role_line: Optional[str] = None
        _section_hdr = re.compile(
            r'^(WORK|EXPERIENCE|EMPLOYMENT|PROFESSIONAL\s+EXPERIENCE|CAREER|HISTORY|EDUCATION|SKILLS|PROJECTS)\b',
            re.IGNORECASE,
        )

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Look for date ranges that indicate work experience
            date_match = None
            start_date = None
            end_date = None

            # Try dateparser-powered extraction first
            if _HAS_DATEPARSER:
                date_result = _extract_date_range(line)
                if date_result:
                    start_date, end_date = date_result

            # Fallback to old regex patterns
            if not start_date:
                for pattern in self.date_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        date_match = match
                        start_date = match.group(1)
                        end_date = match.group(2).lower()
                        if end_date in ['present', 'current', 'now']:
                            end_date = 'present'
                        break

            if start_date:
                # If we have a current job being built, save it
                if current_job and (current_job.get('company') or current_job.get('title')):
                    jobs.append(current_job)

                # Start a new job entry
                # start_date and end_date already set above

                # Try to extract company and title from the same line or previous lines
                if date_match:
                    parts = line[:date_match.start()].strip()
                else:
                    # dateparser path: find split position in original line
                    parts = ""
                    normalized = _normalize_date_text(line)
                    m = _DATE_RANGE_CANDIDATE_RE.search(normalized)
                    if m:
                        raw_start = m.group(1).strip()
                        # Try to locate raw_start in original line (handle period differences)
                        candidates = [raw_start]
                        words = raw_start.split()
                        if len(words) >= 2 and not words[0].endswith('.'):
                            candidates.append(words[0] + '. ' + ' '.join(words[1:]))
                            candidates.append(words[0].upper() + '. ' + ' '.join(words[1:]))
                        for candidate in candidates:
                            idx = line.lower().find(candidate.lower())
                            if idx >= 0:
                                parts = line[:idx].strip()
                                break
                        if not parts:
                            # Fallback: find year and walk back for month
                            year_match = re.search(r'\d{4}', raw_start)
                            if year_match:
                                year = year_match.group(0)
                                for ym in re.finditer(r'\b' + re.escape(year) + r'\b', line):
                                    before = line[:ym.start()]
                                    month_pat = re.search(r'\b([A-Za-z]{3,}\.?)\s+$', before)
                                    if month_pat:
                                        parts = line[:month_pat.start()].strip()
                                        break
                                    parts = line[:ym.start()].strip()
                                    break
                if not parts and last_role_line:
                    parts = last_role_line
                last_role_line = None

                title, company = self._split_title_company(parts)

                current_job = {
                    'company': company,
                    'title': title,
                    'start_date': start_date,
                    'end_date': end_date,
                    'description': ''
                }
            elif current_job is not None:
                # Accumulate job description
                if len(line) > 20:  # Likely a description line
                    current_job['description'] += line + ' '
            else:
                # Between jobs: remember short lines as possible role/company headers
                if len(line) < 120 and not _section_hdr.match(line) and not line.startswith(('-', '•', '*', '·', '▸')):
                    last_role_line = line

        # Don't forget the last job
        if current_job and (current_job.get('company') or current_job.get('title')):
            jobs.append(current_job)

        return jobs

    # Broad list used for full-text fallback when no skills section is found
    KNOWN_SKILLS_BROAD = [
        # Languages
        "python", "java", "javascript", "typescript", "c++", "c#", "c", "golang",
        "rust", "scala", "kotlin", "swift", "ruby", "php", "r", "matlab", "ada",
        "assembly", "bash", "powershell", "perl",
        # Web / frameworks
        "react", "vue", "angular", "node", "django", "flask", "fastapi", "spring",
        ".net", "qt", "boost", "opencv", "grpc",
        # Databases
        "sql", "postgresql", "mysql", "mongodb", "redis", "oracle", "sqlite",
        "elasticsearch", "cassandra", "dynamodb",
        # Cloud / DevOps
        "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
        "jenkins", "git", "linux", "unix", "nginx", "ci/cd", "devops",
        # Architecture / design
        "uml", "ooad", "design patterns", "microservices", "soa", "rest",
        "system design", "software architecture", "requirement engineering",
        # Embedded / systems
        "embedded", "rtos", "qnx", "vxworks", "freertos", "embedded linux",
        "microcontroller", "fpga", "arm", "tcp/ip", "can bus", "modbus",
        "uart", "spi", "i2c", "ipc", "multithreading", "firmware", "bsp",
        "device driver", "bootloader", "real-time",
        # Safety standards
        "sil4", "do-178", "iso 26262", "misra", "functional safety",
        # Data / AI
        "machine learning", "deep learning", "ai", "nlp", "data analysis",
        "spark", "hadoop", "kafka", "tableau", "power bi", "excel",
        # Project / leadership
        "agile", "scrum", "kanban", "jira", "project management",
        "leadership", "mentoring", "management", "communication",
        # Testing
        "unit testing", "tdd", "cmake", "code review", "ci/cd",
    ]

    def _extract_skills(self, text: str) -> List[str]:
        skills = []

        # ── Step 1: skills section extraction ──────────────────────────────────
        skill_headers = [
            r'SKILLS?\s*:?\s*\n',
            r'TECHNICAL\s+SKILLS?\s*:?\s*\n',
            r'KEY\s+SKILLS?\s*:?\s*\n',
            r'CORE\s+(?:SKILLS?|COMPETENCIES?)\s*:?\s*\n',
            r'COMPETENCIES?\s*:?\s*\n',
            r'EXPERTISE\s*:?\s*\n',
            r'TECHNICAL\s+EXPERTISE\s*:?\s*\n',
            r'TECHNOLOGIES?\s*:?\s*\n',
            r'TECH(?:NICAL)?\s+STACK\s*:?\s*\n',
            r'PROGRAMMING\s+LANGUAGES?\s*:?\s*\n',
            r'TOOLS\s+(?:AND|&)?\s*TECHNOLOGIES?\s*:?\s*\n',
            r'PROFICIENCIES?\s*:?\s*\n',
            r'CAPABILITIES?\s*:?\s*\n',
            r'AREAS?\s+OF\s+EXPERTISE\s*:?\s*\n',
            r'TECHNICAL\s+PROFICIENCY\s*:?\s*\n',
        ]
        skills_section = ""
        for header in skill_headers:
            match = re.search(
                header + r'(.*?)(?:\n\n|\n[A-Z][A-Z\s]+\n|$)',
                text, re.DOTALL | re.IGNORECASE
            )
            if match:
                skills_section = match.group(1)
                break

        if skills_section:
            # Split on common separators including " and " / " & " conjunctions
            raw_skills = re.split(r'[,;|•\-\n]|(?:\s+and\s+)|(?:\s+&\s+)', skills_section)
            skills = [s.strip() for s in raw_skills if 0 < len(s.strip()) < 60]

        # ── Step 2: full-text scan with flashtext (MASTER_SKILLS) ────────────
        try:
            from app.backend.services.skill_matcher import _extract_skills_from_text
            scanned = _extract_skills_from_text(text)
            # Merge — prefer section-extracted, then add scanned extras
            existing = {s.lower() for s in skills}
            skills.extend(s for s in scanned if s.lower() not in existing)
        except Exception as e:
            # Final fallback: original KNOWN_SKILLS_BROAD list
            logger.warning("Skills registry unavailable, using fallback KNOWN_SKILLS_BROAD list: %s", e)
            if not skills:
                text_lower = text.lower()
                skills.extend(s for s in self.KNOWN_SKILLS_BROAD if s in text_lower)

        # ── Step 3: Normalize aliases and deduplicate ─────────────────────────
        try:
            from app.backend.services.skill_matcher import normalize_skill_name
            normalized = []
            seen = set()
            for skill in skills:
                canonical = normalize_skill_name(skill)
                key = canonical.lower()
                if key not in seen:
                    seen.add(key)
                    normalized.append(canonical)
            skills = normalized
        except Exception:
            pass  # normalization is best-effort; dedup still applies

        return list(dict.fromkeys(skills))  # final dedup, preserve order

    # ── Education extraction helpers ────────────────────────────────────
    _DEGREE_PATTERNS = re.compile(
        r'\b('
        # Doctoral
        r'Ph\.?D\.?|Doctorate|D\.?Phil\.?|Ed\.?D\.?|'
        # Master's
        r"Master(?:'?s)?|M\.?S\.?c?\.?|M\.?A\.?|M\.?B\.?A\.?|M\.?E\.?|M\.?Tech\.?|"
        r'M\.?Phil\.?|M\.?Ed\.?|M\.?Com\.?|M\.?Sc\.?|MCA|'
        # Bachelor's
        r"Bachelor(?:'?s)?|B\.?S\.?c?\.?|B\.?A\.?|B\.?E\.?|B\.?Tech\.?|"
        r'B\.?Com\.?|B\.?Sc\.?|BCA|B\.?B\.?A\.?|B\.?Arch\.?|'
        # Associate / Diploma / Certificate
        r"Associate(?:'?s)?|Diploma|Advanced\s+Diploma|National\s+Diploma|"
        r'Certificate|Professional\s+Certificate|'
        # International
        r'Ingenieur|Dipl\.?\s*Ing\.?|Licence|Laurea|Magister|'
        # Secondary / High School
        r'Higher\s+Secondary(?:\s+Education)?|Secondary(?:\s+Education)?|'
        r'High\s+School(?:\s+Diploma)?|HSC|SSC|SSLC|GED|A[\s-]?Levels?|O[\s-]?Levels?|'
        # Postgraduate generic
        r'Post\s*Graduate|Postgraduate|PG\s+Diploma'
        r')\b',
        re.IGNORECASE
    )

    _INSTITUTION_SUFFIXES = re.compile(
        r'\b[A-Z][A-Za-z\s&\'-]+'
        r'(?:University|College|Institute|School|Academy|Polytechnic|'
        r'Conservatory|Seminary|Centre|Center|Foundation)'
        r'(?:\s+of\s+[A-Z][A-Za-z\s&\'-]+?)?\b',
        re.IGNORECASE
    )

    _KNOWN_INSTITUTIONS = {
        'MIT', 'IIT', 'NIT', 'IIIT', 'BITS', 'ISB', 'IIM',
        'Yale', 'Harvard', 'Stanford', 'Princeton', 'Oxford', 'Cambridge',
        'LSE', 'ETH', 'Caltech', 'UCLA', 'USC', 'NYU', 'CMU',
    }

    # Qualifiers that are part of the degree name, not the field
    _DEGREE_QUALIFIERS = (
        r'(?:Engineering|Science|Arts|Commerce|Technology|'
        r'Business\s+Administration|Computer\s+Applications|'
        r'Law|Medicine|Education|Philosophy|Fine\s+Arts|'
        r'Social\s+Work|Public\s+Health|Public\s+Administration|'
        r'Divinity|Architecture|Design|Music|Nursing|Pharmacy|Agriculture)'
    )

    # Degree words used in "degree of/in" patterns (mirrors _DEGREE_PATTERNS)
    _DEGREE_WORDS = (
        r"(?:Ph\.?D\.?|Doctorate|D\.?Phil\.?|Ed\.?D\.?|"
        r"Master(?:'?s)?|M\.?S\.?c?\.?|M\.?A\.?|M\.?B\.?A\.?|M\.?E\.?|M\.?Tech\.?|"
        r"M\.?Phil\.?|M\.?Ed\.?|M\.?Com\.?|M\.?Sc\.?|MCA|"
        r"Bachelor(?:'?s)?|B\.?S\.?c?\.?|B\.?A\.?|B\.?E\.?|B\.?Tech\.?|"
        r"B\.?Com\.?|B\.?Sc\.?|BCA|B\.?B\.?A\.?|B\.?Arch\.?|"
        r"Associate(?:'?s)?|Diploma|Advanced\s+Diploma|National\s+Diploma|"
        r"Certificate|Professional\s+Certificate|"
        r"Higher\s+Secondary(?:\s+Education)?|Secondary(?:\s+Education)?|"
        r"High\s+School(?:\s+Diploma)?|HSC|SSC|SSLC|GED|"
        r"Post\s*Graduate|Postgraduate|PG\s+Diploma)"
    )

    # Ending boundary for field-of-study captures
    _FIELD_BOUNDARY = r"(?:\s+(?:from|at)\b|\s*[,\u2013\u2014\-]\s*|\s+\d{4}|\s*$)"

    def _extract_field_of_study(self, line: str, degree_match) -> str:
        """Extract field of study from education line.

        E.g. 'Bachelor of Engineering Electronics and Instrumentation'
        -> 'Electronics and Instrumentation'
        """
        # Pattern 1: "Bachelor/Master of QUALIFIER [in] FIELD" (qualifier is part of degree)
        # e.g. "Bachelor of Engineering Electronics and Instrumentation"
        # e.g. "Master of Science in Computer Science"
        m = re.search(
            r"(?:Bachelor|Master)(?:\s*(?:'s))?\s+of\s+" + self._DEGREE_QUALIFIERS +
            r"(?:\s+in)?\s+(.+?)(?:" + self._FIELD_BOUNDARY + r")",
            line, re.IGNORECASE)
        if m:
            field = m.group(1).strip()
            field = re.sub(r'\s+(?:University|College|Institute|School|Academy).*$', '', field, flags=re.IGNORECASE)
            return field

        # Pattern 2: "degree [of QUALIFIER] in FIELD"
        # e.g. "B.A. in Economics", "Ph.D. in Data Science"
        m = re.search(
            self._DEGREE_WORDS + r"(?:\s*(?:'s))?(?:\s+of\s+" + self._DEGREE_QUALIFIERS + r")?\s+in\s+(.+?)(?:" + self._FIELD_BOUNDARY + r")",
            line, re.IGNORECASE)
        if m:
            field = m.group(1).strip()
            field = re.sub(r'\s+(?:University|College|Institute|School|Academy).*$', '', field, flags=re.IGNORECASE)
            return field

        # Pattern 3: "degree of FIELD" (direct, no qualifier)
        # e.g. "Bachelor of Commerce", "Diploma of Accounting"
        m = re.search(
            self._DEGREE_WORDS + r"(?:\s*(?:'s))?\s+of\s+(.+?)(?:" + self._FIELD_BOUNDARY + r")",
            line, re.IGNORECASE)
        if m:
            field = m.group(1).strip()
            field = re.sub(r'\s+(?:University|College|Institute|School|Academy).*$', '', field, flags=re.IGNORECASE)
            return field

        # Pattern 4: Degree followed directly by field text (no preposition)
        # e.g. "Higher Secondary Education Physics, Chemistry, Math and Computer Science 1996"
        # e.g. "B.E. Electronics and Instrumentation 1998"
        degree_end = degree_match.end()
        remainder = line[degree_end:].strip()
        # Skip trailing dots/commas/periods that weren't consumed by the degree regex
        remainder = re.sub(r'^[.,;:\s]+', '', remainder)
        # Remove year(s) from the remainder
        remainder_no_year = re.sub(r'\b(19|20)\d{2}\b.*$', '', remainder).strip()
        # Remove trailing dashes/dots/commas
        remainder_no_year = re.sub(r'[\s\u2013\u2014\-–—]+$', '', remainder_no_year).strip()
        # Remove "from/at INSTITUTION" patterns (including at start of remainder)
        remainder_no_year = re.sub(r'(?:^|\s+)(?:at|from)\s+.*$', '', remainder_no_year, flags=re.IGNORECASE).strip()
        # Remove institution suffix patterns (e.g. "Harvard Business School")
        remainder_no_year = re.sub(
            r',?\s*[A-Z][A-Za-z\s&\'-]+'
            r'(?:University|College|Institute|School|Academy|Polytechnic|'
            r'Centre|Center|Foundation).*$',
            '', remainder_no_year, flags=re.IGNORECASE).strip()
        # Remove known abbreviated institution names
        for inst in self._KNOWN_INSTITUTIONS:
            remainder_no_year = re.sub(
                r',?\s*\b' + re.escape(inst) + r'\b.*$', '', remainder_no_year, flags=re.IGNORECASE).strip()
        # If what's left is substantive (>=3 chars), use it as the field
        if len(remainder_no_year) >= 3:
            return remainder_no_year

        # Pattern 5: Common field keywords on same line as degree
        field_keywords = re.compile(
            r'\b(Computer\s+Science|Electrical\s+Engineering|Mechanical\s+Engineering|'
            r'Civil\s+Engineering|Electronics(?:\s+and\s+\w+)*|Information\s+Technology|'
            r'Business\s+Administration|Commerce|Mathematics|Physics|Chemistry|'
            r'Biology|Economics|Accounting|Finance|Marketing|'
            r'Software\s+Engineering|Data\s+Science|Artificial\s+Intelligence|'
            r'Arts|Science|Engineering|Instrumentation|Communication)\b',
            re.IGNORECASE)
        m = field_keywords.search(line)
        if m:
            return m.group(1).strip()

        return ""

    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        education = []

        # Look for education section
        edu_headers = [r'EDUCATION\s*:?\s*\n', r'ACADEMIC\s+BACKGROUND\s*:?\s*\n', r'QUALIFICATIONS\s*:?\s*\n']

        edu_section = ""
        for header in edu_headers:
            match = re.search(header + r'(.*?)(?:\n\n[A-Z][A-Z\s]+\n|\Z)', text, re.DOTALL | re.IGNORECASE)
            if match:
                edu_section = match.group(1)
                break

        if edu_section:
            lines = edu_section.strip().split('\n')
            for line in lines:
                line = line.strip()
                if len(line) < 10:
                    continue

                # Look for degree patterns (broad coverage)
                degree_match = self._DEGREE_PATTERNS.search(line)
                if degree_match:
                    university = ""
                    field = ""
                    year = ""

                    # Try to extract year
                    year_match = re.search(r'\b(19|20)\d{2}\b', line)
                    if year_match:
                        year = year_match.group(0)

                    # Try to extract university (suffix-based)
                    uni_match = self._INSTITUTION_SUFFIXES.search(line)
                    if uni_match:
                        university = uni_match.group(0).strip()
                    else:
                        # Try "from/at" prefix pattern
                        uni_match = re.search(r'(?:from\s+|at\s+)([A-Z][A-Za-z\s&\'-]+(?:University|College|Institute|School|Academy|Polytechnic|Centre|Center|Foundation))', line, re.IGNORECASE)
                        if uni_match:
                            university = uni_match.group(1).strip()

                    # Also check for known abbreviated institutions on this line
                    if not university:
                        words_on_line = set(re.findall(r'\b[A-Za-z]+\b', line))
                        for inst in self._KNOWN_INSTITUTIONS:
                            if inst in words_on_line or inst.upper() in {w.upper() for w in words_on_line}:
                                university = inst
                                break

                    # Extract field of study
                    field = self._extract_field_of_study(line, degree_match)

                    education.append({
                        'degree': degree_match.group(0),
                        'field': field,
                        'university': university,
                        'year': year
                    })

        return education

    def _extract_certifications(self, text: str) -> list:
        """Extract certifications from resume text."""
        certifications = []
        lines = text.split('\n')
        in_cert_section = False

        cert_headers = re.compile(
            r'^\s*(CERTIFICATIONS?|CERTIFICATES?|PROFESSIONAL\s+CERTIFICATIONS?|'
            r'LICENSES?\s*(?:AND|&)?\s*CERTIFICATIONS?|ACCREDITATIONS?)\s*:?\s*$',
            re.IGNORECASE)

        # Generic section header to detect end of cert section
        section_header = re.compile(
            r'^\s*[A-Z][A-Z\s&/]{2,}(?:\s*:)?\s*$')

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if cert_headers.match(stripped):
                in_cert_section = True
                continue

            if in_cert_section:
                # Check if we hit the next section
                if section_header.match(stripped) and not cert_headers.match(stripped):
                    in_cert_section = False
                    continue
                # Extract certification (strip bullet points, dashes, numbers)
                cert = re.sub(r'^[\s\-•●○◆▪►\d.)\]]+\s*', '', stripped)
                if cert and len(cert) > 3:  # Skip very short lines
                    certifications.append(cert)

        return certifications

    def _extract_languages(self, text: str) -> list:
        """Extract languages from resume text."""
        languages = []
        lines = text.split('\n')
        in_lang_section = False

        lang_headers = re.compile(
            r'^\s*(LANGUAGES?|LANGUAGE\s+SKILLS?|LANGUAGE\s+PROFICIENCY|'
            r'LINGUISTIC\s+SKILLS?)\s*:?\s*$',
            re.IGNORECASE)

        section_header = re.compile(
            r'^\s*[A-Z][A-Z\s&/]{2,}(?:\s*:)?\s*$')

        proficiency_re = re.compile(
            r'[-–—:]\s*(native|fluent|proficient|intermediate|'
            r'advanced|basic|beginner|elementary|conversational|'
            r'professional|working\s+proficiency|mother\s+tongue|'
            r'native\s+speaker|bilingual)\b',
            re.IGNORECASE)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if lang_headers.match(stripped):
                in_lang_section = True
                continue

            if in_lang_section:
                if section_header.match(stripped) and not lang_headers.match(stripped):
                    in_lang_section = False
                    continue

                # Could be comma-separated on one line or one per line
                # Try splitting by comma first
                items = re.split(r'[,;]', stripped)
                for item in items:
                    item = re.sub(r'^[\s\-•●○◆▪►\d.)\]]+\s*', '', item).strip()
                    if not item or len(item) < 2:
                        continue

                    # Check for proficiency level
                    prof_match = proficiency_re.search(item)
                    if prof_match:
                        lang_name = item[:prof_match.start()].strip().rstrip('-–—: ')
                        proficiency = prof_match.group(1).strip().capitalize()
                    else:
                        # Check for parenthesized proficiency
                        paren_match = re.search(r'\(([^)]+)\)', item)
                        if paren_match:
                            lang_name = item[:paren_match.start()].strip()
                            proficiency = paren_match.group(1).strip().capitalize()
                        else:
                            lang_name = item.strip()
                            proficiency = ""

                    if lang_name and len(lang_name) > 1:
                        languages.append({
                            "language": lang_name,
                            "proficiency": proficiency
                        })

        return languages

    def _extract_professional_summary(self, text: str) -> str:
        """Extract professional summary/objective from resume text."""
        lines = text.split('\n')
        in_summary_section = False
        summary_lines = []

        summary_headers = re.compile(
            r'^\s*(PROFESSIONAL\s+SUMMARY|SUMMARY|CAREER\s+SUMMARY|'
            r'EXECUTIVE\s+SUMMARY|OBJECTIVE|CAREER\s+OBJECTIVE|'
            r'PROFILE|PROFESSIONAL\s+PROFILE|ABOUT\s+ME|OVERVIEW)\s*:?\s*$',
            re.IGNORECASE)

        section_header = re.compile(
            r'^\s*[A-Z][A-Z\s&/]{2,}(?:\s*:)?\s*$')

        for line in lines:
            stripped = line.strip()

            if summary_headers.match(stripped):
                in_summary_section = True
                continue

            if in_summary_section:
                if not stripped:
                    # Allow one blank line within summary
                    if summary_lines and summary_lines[-1] != '':
                        summary_lines.append('')
                    continue

                if section_header.match(stripped) and not summary_headers.match(stripped):
                    in_summary_section = False
                    break

                summary_lines.append(stripped)

        # Join and truncate to 500 chars
        summary = ' '.join(line for line in summary_lines if line).strip()
        return summary[:500] if summary else ""

    def _extract_name(self, text: str) -> str:
        """
        Extract candidate name from the top of the resume.
        Names are typically the first prominent line: 1-5 capitalised words,
        not a common section header.

        Note: older logic rejected any line containing a hyphen in a character class,
        which incorrectly skipped hyphenated names (e.g. Mary-Jane Smith). We now
        split on | / • and treat phone/email segments separately.
        """
        SKIP_WORDS = {
            'resume', 'curriculum', 'vitae', 'cv', 'profile', 'summary',
            'objective', 'contact', 'address', 'details', 'information',
            'page', 'updated', 'date', 'domain', 'experience', 'education',
            'skills', 'employment', 'work', 'projects', 'references',
            'certifications', 'awards', 'publications', 'languages',
            'interests', 'hobbies', 'activities', 'achievements',
        }
        lines = text.strip().split('\n')
        for line in lines[:15]:
            line = line.strip()
            if not line or len(line) < 3:
                continue
            name = self._name_from_header_line(line, SKIP_WORDS)
            if name:
                return name
        return ''

    def _name_from_header_line(self, line: str, skip_words: set) -> str:
        """Try to read a person name from one header line (possibly 'Name | phone | city')."""
        # Additional phrases that are definitely not names
        skip_phrases = {
            'software development', 'software engineer', 'software engineering',
            'data scientist', 'data science', 'machine learning', 'deep learning',
            'full stack', 'frontend developer', 'backend developer', 'web developer',
            'mobile developer', 'devops engineer', 'systems engineer',
            'project manager', 'product manager', 'scrum master', 'agile coach',
            'business analyst', 'quality assurance', 'qa engineer', 'test engineer',
            'database administrator', 'network engineer', 'security engineer',
            'cloud architect', 'solution architect', 'technical lead', 'team lead',
            'engineering manager', 'cto', 'cio', 'ceo', 'cfo', 'coo',
            'vice president', 'director of', 'head of', 'lead of',
            # Section headers that are commonly mistaken for names
            'key expertise', 'key skills', 'core competencies', 'top skills',
            'technical experience', 'embedded experience', 'professional experience',
            'work experience', 'career highlights', 'summary', 'objective',
            'personal details', 'contact information', 'contact details',
        }
        segments = re.split(r'\s*[|•]\s*', line)
        for seg in segments:
            seg = seg.strip()
            if not seg or len(seg) < 2:
                continue
            if '@' in seg or re.search(r'linkedin\.com/', seg, re.IGNORECASE):
                continue
            if sum(1 for c in seg if c.isdigit()) > 2:
                continue
            if re.search(r'\+?\d[\d\s().\-]{8,}\d', seg):
                continue
            words = seg.split()
            if not (1 <= len(words) <= 5):
                continue
            # Check individual words against skip_words
            if any(w.lower() in skip_words for w in words):
                continue
            # Check entire phrase against skip_phrases
            seg_lower = seg.lower()
            if any(phrase in seg_lower for phrase in skip_phrases):
                continue
            # Check for job title patterns (e.g., "Developer" at end)
            if any(w.lower() in ('developer', 'engineer', 'manager', 'analyst', 'architect', 'lead', 'consultant') for w in words[-1:]):
                continue
            # Name should have Title Case (most words capitalized)
            cap_count = sum(1 for w in words if w and w[0].isupper())
            if cap_count >= max(1, len(words) - 1):
                return seg
        return ''

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        info = {}

        # Name — extracted first so we can use it in the result
        name = self._extract_name(text)
        if name:
            info['name'] = name

        # Email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            info['email'] = email_match.group(0)

        # Phone - stricter pattern to avoid matching years
        # Must have at least 7 digits and look like a real phone number
        phone_patterns = [
            # International: +1-555-123-4567, +91 98765 43210
            r'\+[0-9]{1,3}[-\s.]?[0-9]{1,4}[-\s.]?[0-9]{1,4}[-\s.]?[0-9]{1,9}',
            # With parentheses: (555) 123-4567, (555)123-4567
            r'\([0-9]{2,4}\)\s*[0-9]{1,4}[-\s.]?[0-9]{1,9}',
            # Standard: 555-123-4567, 555.123.4567, 555 123 4567
            r'\b[0-9]{3,4}[-\s.][0-9]{3,4}[-\s.][0-9]{4}\b',
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                phone = phone_match.group(0).strip()
                # Validate: must have at least 7 digits and not be a year (1900-2099)
                digits_only = re.sub(r'\D', '', phone)
                if len(digits_only) >= 7:
                    year_check = int(digits_only[:4]) if len(digits_only) >= 4 else 0
                    if not (1900 <= year_check <= 2099):
                        info['phone'] = phone
                        break

        # LinkedIn
        linkedin_match = re.search(r'linkedin\.com/in/[A-Za-z0-9\-_%]+', text, re.IGNORECASE)
        if linkedin_match:
            info['linkedin'] = linkedin_match.group(0)

        return info


def _name_from_email(email: str) -> str:
    """john.doe@x.com → 'John Doe' when local part has two+ alpha tokens."""
    if not email or "@" not in email:
        return ""
    local = email.split("@", 1)[0].strip().lower()
    if not local:
        return ""
    tokens = [t for t in re.split(r"[._+\-]+", local) if t.isalpha() and len(t) >= 2]
    if len(tokens) < 2:
        return ""
    return " ".join(t.capitalize() for t in tokens[:4])


def _extract_name_relaxed(text: str) -> str:
    """
    Fallback when strict header heuristics miss (e.g. title line before name, odd layout).
    """
    skip = {
        "resume", "curriculum", "vitae", "cv", "profile", "summary",
        "objective", "contact", "address", "details", "information",
        "page", "updated", "date", "experience", "education", "skills",
        "employment", "work", "projects", "references",
    }
    title_case_name = re.compile(r"^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\s*$")
    for line in text.strip().split("\n")[:35]:
        line = line.strip()
        if len(line) < 4 or len(line) > 80:
            continue
        if "@" in line or re.search(r"linkedin\.com/", line, re.IGNORECASE):
            continue
        if re.search(r"\+?\d[\d\s().\-]{8,}\d", line):
            continue
        if title_case_name.match(line):
            words = line.split()
            if any(w.lower() in skip for w in words):
                continue
            return line
    return ""


async def _llm_extract_structured(raw_text: str, data: Dict[str, Any]) -> Dict[str, Any]:
    """LLM fallback for structured field gaps when deterministic extraction misses entries.

    Checks heuristic gap indicators; if nothing is missing, returns empty dict immediately.
    Otherwise sends first 4000 chars of resume to the LLM with a focused prompt,
    requesting only the missing fields as JSON.
    Graceful 10 s timeout — if LLM is unavailable, returns empty dict.
    """
    import json as _json
    import httpx as _httpx

    OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "gemma4:31b-cloud")

    # --- Heuristic gap detection ---
    needs_work = False
    needs_education = False

    work_experience = data.get("work_experience", [])
    education = data.get("education", [])

    # Work gap: ≤1 entry but text has 2+ company-keyword mentions
    if len(work_experience) <= 1 and raw_text:
        company_keywords = re.findall(
            r"\b(?:Ltd|Inc|Corp|Pvt|LLC|Co\.|Corporation|Company|Technologies|Solutions|Services|Enterprises|Group|Industries|Consulting|Associates|Systems|Labs|Ventures|Holdings)\b",
            raw_text,
            re.IGNORECASE,
        )
        if len(company_keywords) >= 2:
            needs_work = True

    # Education gap: no entries but text mentions degree keywords
    if len(education) == 0 and raw_text:
        degree_keywords = re.findall(
            r"\b(?:Bachelor|Master|B\.?Tech|M\.?Tech|B\.?E|M\.?E|B\.?Sc|M\.?Sc|MBA|PhD|BBA|BCA|MCA|Diploma|Associate\s+Degree|Doctorate|B\.?Com|M\.?Com|B\.?A|M\.?A|B\.?Arch)\b",
            raw_text,
            re.IGNORECASE,
        )
        if degree_keywords:
            needs_education = True

    if not needs_work and not needs_education:
        return {}

    # --- Build focused prompt ---
    missing_fields: list[str] = []
    if needs_work:
        missing_fields.append("work_experience")
    if needs_education:
        missing_fields.append("education")

    system_prompt = (
        "You are a resume data extractor. Extract ONLY the requested fields "
        "and return ONLY valid JSON. Be precise and conservative."
    )

    fields_desc = []
    if needs_work:
        fields_desc.append(
            '"work_experience": list of objects with keys "title", "company", '
            '"start_date", "end_date", "description". Use null for missing sub-fields.'
        )
    if needs_education:
        fields_desc.append(
            '"education": list of objects with keys "degree", "institution", '
            '"start_date", "end_date", "gpa". Use null for missing sub-fields.'
        )

    user_prompt = (
        f"Extract the following fields from this resume: {', '.join(missing_fields)}.\n"
        f"Return ONLY a valid JSON object with these keys:\n"
        + "\n".join(f"  - {desc}" for desc in fields_desc)
        + "\n\nResume text:\n"
        + raw_text[:4000]
        + "\n\nJSON output:"
    )

    try:
        from app.backend.services.llm_service import get_ollama_headers

        headers = get_ollama_headers(OLLAMA_BASE_URL)

        async with _httpx.AsyncClient(timeout=10.0) as client:
            response = await client.post(
                f"{OLLAMA_BASE_URL}/api/chat",
                headers=headers,
                json={
                    "model": OLLAMA_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "stream": False,
                    "options": {
                        "temperature": 0.1,
                        "num_predict": 1500,
                    },
                },
            )
            response.raise_for_status()

            result = response.json()
            llm_output = result.get("message", {}).get("content", "").strip()

            # Strip markdown code blocks
            if "```json" in llm_output:
                llm_output = llm_output.split("```json")[1].split("```")[0].strip()
            elif "```" in llm_output:
                llm_output = llm_output.split("```")[1].split("```")[0].strip()

            # Minimum length check for safety
            if len(llm_output) < 10:
                logger.warning("[LLM Structured] Response too short (%d chars), discarding", len(llm_output))
                return {}

            parsed = _json.loads(llm_output)

            if not isinstance(parsed, dict):
                logger.warning("[LLM Structured] Non-dict response: %s", type(parsed).__name__)
                return {}

            # Keep only requested fields
            return {k: v for k, v in parsed.items() if k in missing_fields and isinstance(v, list)}

    except _json.JSONDecodeError as e:
        logger.warning("[LLM Structured] JSON parse error: %s", e)
        return {}
    except _httpx.TimeoutException:
        logger.warning("[LLM Structured] Timed out after 10s")
        return {}
    except Exception as e:
        logger.warning("[LLM Structured] Failed: %s: %s", type(e).__name__, str(e)[:200])
        return {}


async def enrich_parsed_resume_async(data: Dict[str, Any], filename: Optional[str] = None) -> None:
    """Fill gaps in parser output in-place using LLM + NER + regex fallbacks."""
    from app.backend.services.llm_contact_extractor import extract_contact_with_llm, merge_contact_info
    
    contact = data.setdefault("contact_info", {})
    raw = (data.get("raw_text") or "").strip()
    
    # --- Contact enrichment (skip if already complete) ---
    if not (contact.get("name") and contact.get("email")):
        # Try LLM extraction first (most accurate, handles all edge cases)
        llm_contact = None
        if raw:
            try:
                llm_contact = await extract_contact_with_llm(raw, timeout=8.0)
            except Exception as e:
                logger.warning("LLM contact extraction failed, using fallbacks: %s", e)
        
        # Merge LLM results with existing regex results
        if llm_contact:
            contact.update(merge_contact_info(contact, llm_contact))
        
        # If still no name, try traditional fallbacks
        if not contact.get("name"):
            guess = None
            
            # Tier 1: Try spaCy NER
            if raw:
                guess = _extract_name_ner(raw)
            
            # Tier 2: Fallback to email-based extraction
            if not guess:
                email = (contact.get("email") or "").strip()
                guess = _name_from_email(email)
            
            # Tier 3: Fallback to relaxed header scan
            if not guess and raw:
                guess = _extract_name_relaxed(raw)
            
            # Tier 4: Fallback to filename-based extraction
            if not guess and filename:
                guess = _name_from_filename(filename)
            
            if guess:
                contact["name"] = guess
    
    # --- LLM fallback for structured field gaps ---
    try:
        llm_fields = await _llm_extract_structured(
            data.get("raw_text", ""), data)
        
        # Merge: only fill gaps, don't overwrite deterministic results
        if "work_experience" in llm_fields:
            llm_jobs = llm_fields["work_experience"]
            if isinstance(llm_jobs, list) and len(llm_jobs) > len(data.get("work_experience", [])):
                data["work_experience"] = llm_jobs
        
        if "education" in llm_fields:
            llm_edu = llm_fields["education"]
            if isinstance(llm_edu, list) and len(llm_edu) > len(data.get("education", [])):
                data["education"] = llm_edu
    except Exception:
        pass  # Graceful failure


def enrich_parsed_resume(data: Dict[str, Any], filename: Optional[str] = None) -> None:
    """Synchronous wrapper - Fill gaps in parser output in-place (name from NER / email / relaxed header scan / filename)."""
    contact = data.setdefault("contact_info", {})
    raw = (data.get("raw_text") or "").strip()
    if (contact.get("name") or "").strip():
        return
    
    # Tier 0: Try spaCy NER first (most accurate for diverse formats)
    guess = None
    if raw:
        guess = _extract_name_ner(raw)
    
    # Tier 1: Fallback to email-based extraction
    if not guess:
        email = (contact.get("email") or "").strip()
        guess = _name_from_email(email)
    
    # Tier 2: Final fallback to relaxed header scan
    if not guess and raw:
        guess = _extract_name_relaxed(raw)
    
    # Tier 3: Fallback to filename-based extraction
    if not guess and filename:
        guess = _name_from_filename(filename)

    if guess:
        contact["name"] = guess


def _name_from_filename(filename: str) -> str:
    """Extract a name from the resume filename.
    
    Examples:
        "Suhas Mullangi.pdf" → "Suhas Mullangi"
        "john_doe_resume_2024.pdf" → "John Doe"
        "resume_jane_smith.docx" → "Jane Smith"
    """
    import re
    
    if not filename:
        return ""
    
    # Strip extension
    name = filename.rsplit(".", 1)[0] if "." in filename else filename
    
    # First, replace underscores and hyphens with spaces to enable word boundary matching
    name = re.sub(r"[_\-]+", " ", name)
    
    # Remove common non-name patterns (case-insensitive)
    # Patterns: resume, cv, curriculum, vitae, updated, final, latest, revised, new, copy
    name = re.sub(r"(?i)\b(resume|cv|curriculum|vitae|updated?|final|latest|revised?|new|copy)\b", "", name)
    
    # Remove dates in various formats (YYYY, YY, MM-DD, DD-MM, etc.)
    name = re.sub(r"\b(19|20)\d{2}\b", "", name)  # Years like 2024, 1995
    name = re.sub(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", "", name)  # Date patterns
    name = re.sub(r"\b\d{4}\b", "", name)  # 4-digit numbers (years)
    
    # Clean up multiple spaces and strip
    name = re.sub(r"\s+", " ", name).strip()
    
    # Title case the result
    name = name.title()
    
    # Validate: must have 2-5 words and no digits
    words = name.split()
    if len(words) < 2 or len(words) > 5:
        return ""
    if any(char.isdigit() for char in name):
        return ""
    
    return name


def parse_resume(file_bytes: bytes, filename: str) -> Dict[str, Any]:
    start = time.monotonic()
    parser = ResumeParser()
    out = parser.parse_resume(file_bytes, filename)
    enrich_parsed_resume(out, filename)
    RESUME_PARSE_DURATION.observe(time.monotonic() - start)
    return out
